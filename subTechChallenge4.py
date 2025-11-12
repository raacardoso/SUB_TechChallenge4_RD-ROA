import cv2
from ultralytics import YOLO
from docx import Document
from docx.shared import Inches
from collections import defaultdict
import os
import numpy as np

# CaminhosSave 
VIDEO_PATH = "camera_tailândia_2.mp4"
OUTPUT_DOC = "Relatorio_Analise_SubTechChallenge4.docx"
THUMBNAILS_DIR = "thumbs"
os.makedirs(THUMBNAILS_DIR, exist_ok=True)

# Modelos YOLO
object_model = YOLO("yolov8m.pt")
pose_model = YOLO("yolov8n-pose.pt")

# ClassesAlvo
target_classes = ["person", "car", "motorcycle", "dog"]
activity_tracking = defaultdict(list)
unusual_events = defaultdict(set)

# ClassificaçãoSev
event_severity = {
    
    "Carro detectado": "Leve",
    "Moto detectada": "Aceitável",
    "Correndo na rua": "Potencial Risco",
    "Cachorro na rua": "Potencial Risco"
}

# Rastreamento
object_seconds = defaultdict(set)
thumbnails_saved = defaultdict(int)
unique_person_ids = set()
unique_dog_ids = set()
unique_motorcycle_ids = set()

cap = cv2.VideoCapture(VIDEO_PATH)
fps = cap.get(cv2.CAP_PROP_FPS)
frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

frame_number = 0
while True:
    ret, frame = cap.read()
    if not ret:
        break

    if frame_number % 5 == 0:
        second = int(frame_number / fps)
        results = object_model.track(frame, persist=True, verbose=False)[0]

        for box in results.boxes:
            cls_id = int(box.cls[0])
            cls_name = object_model.names[cls_id].lower()

            if cls_name in target_classes:
                object_seconds[cls_name].add(second)

                if box.id is not None:
                    obj_id = int(box.id.item())
                    if cls_name == "person":
                        unique_person_ids.add(obj_id)
                    elif cls_name == "dog":
                        unique_dog_ids.add(obj_id)
                    elif cls_name == "motorcycle":
                        unique_motorcycle_ids.add(obj_id)

                if thumbnails_saved[cls_name] < 3:
                    x1, y1, x2, y2 = map(int, box.xyxy[0])
                    cropped = frame[y1:y2, x1:x2]
                    thumb_path = os.path.join(THUMBNAILS_DIR, f"{cls_name}_{second}_{thumbnails_saved[cls_name]}.jpg")
                    cv2.imwrite(thumb_path, cropped)
                    thumbnails_saved[cls_name] += 1

                if cls_name == "dog":
                    unusual_events["Cachorro na rua"].add(second)
                elif cls_name == "car":
                    unusual_events["Carro detectado"].add(second)
                elif cls_name == "motorcycle":
                    unusual_events["Moto detectada"].add(second)

        pose_result = pose_model.predict(frame, verbose=False)[0]
        for kp in pose_result.keypoints:
            person_xy = kp.xy[0].cpu().numpy()
            if person_xy.shape != (17, 2):
                continue

            y_std = np.std(person_xy[:, 1])
            x_std = np.std(person_xy[:, 0])

            if y_std > 90 or x_std > 90:
                action = "Correndo"
            elif y_std > 30 or x_std > 30:
                action = "Andando"
            elif y_std < 10 and x_std < 10:
                action = "Parado"
            else:
                action = "Interagindo"

            activity_tracking[action].append(second)

            if action == "Correndo":
                unusual_events["Correndo na rua"].add(second)

    frame_number += 1

cap.release()

# GeraçãoRelatório
doc = Document()
doc.add_heading(" RD&ROA : Relatório de detecção de riscos Objetos e Atividades", 0)
doc.add_heading(" Ramon C. Sancha RM 356432", 1)
doc.add_heading(" RD&ROA:", 2)
doc.add_paragraph("Este relatório apresenta as detecções de Riscos, objetos e atividades (pessoas, carros, motos, cachorros), classifica eventos incomuns e suas severidades.")

# Seção 1: EstruturaObjetos
for cls in target_classes:
    seconds = sorted(object_seconds[cls])
    doc.add_heading(f"🔹 {cls.capitalize()}", level=1)
    doc.add_paragraph(f"Tempo de detecção: {len(seconds)} segundos.")
    if cls == "person":
        doc.add_paragraph(f"ID's/Quantidade de Pessoas rastreadas: {len(unique_person_ids)}")
    elif cls == "dog":
        doc.add_paragraph(f" ID's/Quantidade de Cachorros rastreados: {len(unique_dog_ids)}")
    elif cls == "motorcycle":
        doc.add_paragraph(f"ID's/Quantidade de Motos rastreadas: {len(unique_motorcycle_ids)}")
    if seconds:
        seconds_str = ', '.join(str(s) + "s" for s in seconds)
        doc.add_paragraph(f" Histórico de detecção: {seconds_str}")
    else:
        doc.add_paragraph(" Nenhuma detecção.")
    for i in range(min(3, thumbnails_saved[cls])):
        if i < len(seconds):
            img_path = os.path.join(THUMBNAILS_DIR, f"{cls}_{seconds[i]}_{i}.jpg")
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(2))

# Seção 2: EstruturaAtividades
doc.add_heading(" - Atividades Humanas Detectadas", level=1)
for activity, seconds in activity_tracking.items():
    unique_secs = sorted(set(seconds))
    doc.add_heading(f" º {activity}", level=2)
    doc.add_paragraph(f"Detectado durante {len(unique_secs)} segundos.")
    doc.add_paragraph(f"HRD - histórico de registro e detecção: {', '.join(str(s)+'s' for s in unique_secs)}")

# Seção 3: EstruturaEventosPoRSeveridade
doc.add_heading("ARCS - Atividades de Risco Classificadas por Severidade", level=1)
severity_levels = ["Leve", "Aceitável", "Potencial Risco", "Perigo"]
total_by_severity = defaultdict(int)

for severity in severity_levels:
    doc.add_heading(f"Severidade: {severity}", level=2)

    for event, level in event_severity.items():
        if level != severity:
            continue

        secs_list = sorted(unusual_events.get(event, []))
        doc.add_heading(f"{event}", level=3)

        if secs_list:
            doc.add_paragraph(f"Ocorreu durante {len(secs_list)} segundos.")
            doc.add_paragraph(f"HRD - histórico de registro e detecção: {', '.join(str(s)+'s' for s in secs_list)}")
            total_by_severity[severity] += len(secs_list)
        else:
            doc.add_paragraph("Nenhuma ocorrência detectada.")

# Totalizador
doc.add_paragraph(" Total por Severidade:")
for sev, count in total_by_severity.items():
    doc.add_paragraph(f"{sev}: total de {count} segundos.")

# Salvar
doc.save(OUTPUT_DOC)
print("Relatório gerado com sucesso:", OUTPUT_DOC)
